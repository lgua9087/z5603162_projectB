"""Build the editable Word report and matching submission PDF from verified artifacts."""

from __future__ import annotations

import argparse
import html
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol

import pandas as pd
import pdfplumber
import pymupdf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image as PillowImage
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.units import mm as rl_mm
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
REPORT_DIR = PROJECT_ROOT / "report"
RESULTS_DIR = PROJECT_ROOT / "results"
DOCX_PATH = REPORT_DIR / "report.docx"
PDF_PATH = REPORT_DIR / "report.pdf"
LIVE_APP_URL = "https://marketready-funds-ustakylve9nywmyokqp4q4.streamlit.app/"
PUBLIC_GITHUB_URL = "https://github.com/lgua9087/marketready-funds"

# Document skill design contract.
DESIGN_PRESET = "narrative_proposal"
HEADER_PATTERN = "editorial_cover"
NAMED_OVERRIDES = {
    "coursework_a4": "A4 portrait; 17 mm side and 16 mm vertical margins",
    "ten_page_narrative": "9.5 pt body, 1.16 spacing, 5 pt paragraph gap",
    "marketready_brand": "navy/teal/gold palette shared with app and figures",
}

NAVY = "16324F"
TEAL = "2A9D8F"
GOLD = "D99B2B"
SLATE = "64748B"
LIGHT = "EEF4F7"
PALE_GOLD = "FFF7E6"
WHITE = "FFFFFF"
BLACK = "1F2933"

FIGURE_NUMBERS = {
    "growth": 1,
    "drawdown": 2,
    "weights": 3,
    "sentiment": 4,
    "fusion": 5,
    "sharpe": 6,
}
TABLE_NUMBERS = {
    "inputs": 1,
    "specifications": 2,
    "performance": 3,
    "combined_risk": 4,
    "sentiment_predictive": 5,
    "fusion": 6,
    "holdings": 7,
    "validation": 8,
    "integrity": 9,
    "annual": 10,
    "cost_fee": 11,
    "fusion_sensitivity": 12,
    "vader": 13,
    "imputation": 14,
}
TOKEN_PATTERN = re.compile(r"\{(fig|table):([a-z_]+)\}")


def rl_hex(value: str) -> colors.Color:
    """Convert the shared six-character palette to a ReportLab colour."""

    return colors.HexColor(f"#{value.lstrip('#')}")


@dataclass(frozen=True)
class ReportData:
    metrics: pd.DataFrame
    weights: pd.DataFrame
    fusion: pd.DataFrame
    vader: pd.DataFrame
    coverage: pd.DataFrame
    validation: pd.DataFrame
    integrity: pd.DataFrame
    combined_risk: pd.DataFrame
    sentiment_predictive: pd.DataFrame
    annual_results: pd.DataFrame
    cost_sensitivity: pd.DataFrame
    fee_sensitivity: pd.DataFrame
    fusion_sensitivity: pd.DataFrame
    imputation: pd.DataFrame
    summary: dict[str, object]


class Builder(Protocol):
    def kicker(self, text: str) -> None: ...

    def title(self, text: str) -> None: ...

    def subtitle(self, text: str) -> None: ...

    def meta(self, text: str) -> None: ...

    def h1(self, text: str) -> None: ...

    def h2(self, text: str) -> None: ...

    def h3(self, text: str) -> None: ...

    def p(self, text: str) -> None: ...

    def callout(self, label: str, text: str) -> None: ...

    def equation(self, text: str) -> None: ...

    def numbered(self, items: list[tuple[str, str]]) -> None: ...

    def table(
        self,
        key: str,
        caption: str,
        headers: list[str],
        rows: list[list[str]],
        widths: list[float],
        *,
        font_size: float = 7.5,
    ) -> None: ...

    def figure(
        self,
        key: str,
        path: Path,
        caption: str,
        *,
        width_inches: float,
        alt_text: str,
    ) -> None: ...

    def new_page(self) -> None: ...


def load_report_data() -> ReportData:
    """Load the final artifact set; no model is recomputed while writing."""

    import json

    tables = RESULTS_DIR / "tables"
    data = RESULTS_DIR / "data"
    metrics = pd.read_csv(tables / "performance_metrics.csv")
    weights = pd.read_csv(data / "fund_weights.csv", parse_dates=["rebalance_date"])
    return ReportData(
        metrics=metrics,
        weights=weights,
        fusion=pd.read_csv(tables / "fusion_comparison.csv"),
        vader=pd.read_csv(tables / "sentiment_model_comparison.csv"),
        coverage=pd.read_csv(tables / "sentiment_coverage_summary.csv"),
        validation=pd.read_csv(tables / "validation_results.csv"),
        integrity=pd.read_csv(tables / "data_integrity_summary.csv"),
        combined_risk=pd.read_csv(tables / "combined_risk_summary.csv"),
        sentiment_predictive=pd.read_csv(
            tables / "sentiment_predictive_diagnostics.csv"
        ),
        annual_results=pd.read_csv(tables / "annual_fund_results.csv"),
        cost_sensitivity=pd.read_csv(tables / "transaction_cost_sensitivity.csv"),
        fee_sensitivity=pd.read_csv(tables / "management_fee_sensitivity.csv"),
        fusion_sensitivity=pd.read_csv(tables / "fusion_parameter_sensitivity.csv"),
        imputation=pd.read_csv(tables / "estimation_imputation_summary.csv"),
        summary=json.loads((data / "report_summary.json").read_text(encoding="utf-8")),
    )


def pct(value: float, digits: int = 1) -> str:
    return f"{value:.{digits}%}"


def number(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def performance_rows(metrics: pd.DataFrame) -> list[list[str]]:
    order = {"Equity": 0, "Crypto": 1, "Combined": 2}
    frame = metrics.assign(_order=metrics["family"].map(order)).sort_values(
        ["_order", "method"]
    )
    return [
        [
            str(row.fund_name),
            pct(float(row.annualized_return)),
            pct(float(row.annualized_volatility)),
            number(float(row.sharpe_ratio)),
            pct(float(row.max_drawdown)),
            pct(float(row.average_turnover)),
        ]
        for row in frame.itertuples()
    ]


def top_holdings_rows(data: ReportData) -> list[list[str]]:
    latest = data.weights.sort_values("rebalance_date").groupby("fund_id").tail(1)
    latest_dates = latest.set_index("fund_id")["rebalance_date"]
    rows: list[list[str]] = []
    names = data.metrics.set_index("fund_id")["fund_name"]
    for fund_id, group in data.weights.groupby("fund_id", sort=False):
        date = latest_dates.loc[fund_id]
        current = group.loc[group["rebalance_date"].eq(date)].nlargest(3, "weight")
        holdings = "; ".join(
            f"{row.ticker} {pct(float(row.weight))}" for row in current.itertuples()
        )
        rows.append([str(names.loc[fund_id]), f"{date:%Y-%m-%d}", holdings])
    return sorted(rows, key=lambda row: row[0])


def combined_risk_rows(data: ReportData) -> list[list[str]]:
    order = {
        "Equal Weight": 0,
        "Minimum Variance": 1,
        "Maximum Sharpe": 2,
        "Risk Parity": 3,
    }
    turnover = data.metrics.set_index("fund_id")["average_turnover"]
    frame = data.combined_risk.assign(
        _order=data.combined_risk["method"].map(order)
    ).sort_values("_order")
    return [
        [
            str(row.method),
            (
                f"{row.median_crypto_capital_weight:.1%} "
                f"({row.minimum_crypto_capital_weight:.1%}-{row.maximum_crypto_capital_weight:.1%})"
            ),
            f"{row.median_crypto_risk_contribution:.1%}",
            f"{row.median_effective_holdings:.1f}",
            f"{row.median_top_5_weight:.1%}",
            f"{turnover.loc[row.fund_id]:.1%}",
        ]
        for row in frame.itertuples()
    ]


def sentiment_predictive_rows(data: ReportData) -> list[list[str]]:
    frame = data.sentiment_predictive.sort_values("horizon_days")
    return [
        [
            f"{int(row.horizon_days)} day" + ("s" if row.horizon_days > 1 else ""),
            f"{row.mean_spearman_ic:+.3f}",
            f"{row.ic_hac_t_stat:+.2f}",
            f"{row.mean_high_minus_low_return * 10_000:+.2f}",
            f"{row.spread_hac_t_stat:+.2f}",
            f"{row.positive_ic_share:.1%}",
        ]
        for row in frame.itertuples()
    ]


def annual_return_rows(data: ReportData) -> list[list[str]]:
    order = {"Equity": 0, "Crypto": 1, "Combined": 2}
    wide = data.annual_results.pivot(
        index=["fund_id", "fund_name", "family", "method"],
        columns="year",
        values="calendar_return",
    ).reset_index()
    wide = wide.assign(_order=wide["family"].map(order)).sort_values(
        ["_order", "method"]
    )
    return [
        [
            str(row["fund_name"]),
            pct(float(row[2021])),
            pct(float(row[2022])),
            pct(float(row[2023])),
        ]
        for _, row in wide.iterrows()
    ]


def cost_fee_rows(data: ReportData) -> list[list[str]]:
    costs = data.cost_sensitivity.loc[
        data.cost_sensitivity["fund_id"].eq("combined_risk_parity")
    ].set_index("transaction_cost_bps")
    fees = data.fee_sensitivity.loc[
        data.fee_sensitivity["fund_id"].eq("combined_risk_parity")
    ].set_index("annual_management_fee")
    scenarios = [
        ("Trading cost: 0 bp", costs.loc[0.0]),
        ("Trading cost: 5 bp; fee: 0%", costs.loc[5.0]),
        ("Trading cost: 10 bp", costs.loc[10.0]),
        ("5 bp cost + 0.5% annual fee", fees.loc[0.005]),
        ("5 bp cost + 1.0% annual fee", fees.loc[0.01]),
    ]
    return [
        [
            label,
            pct(float(row.annualized_return)),
            number(float(row.sharpe_ratio), 3),
            pct(float(row.max_drawdown)),
            f"${row.terminal_wealth:.2f}",
        ]
        for label, row in scenarios
    ]


def fusion_sensitivity_rows(data: ReportData) -> list[list[str]]:
    return [
        [
            str(row.scenario),
            pct(float(row.annualized_return)),
            number(float(row.sharpe_ratio), 3),
            pct(float(row.max_drawdown)),
            pct(float(row.average_turnover)),
        ]
        for row in data.fusion_sensitivity.itertuples()
    ]


def imputation_rows(data: ReportData) -> list[list[str]]:
    return [
        [
            str(row.family),
            f"{int(row.rebalances)}",
            f"{int(row.optimizer_cells):,}",
            f"{int(row.mean_filled_cells):,}",
            pct(float(row.mean_fill_share), digits=2),
            f"{int(row.asset_window_exclusions):,}",
        ]
        for row in data.imputation.itertuples()
    ]


def _replace_static_tokens(text: str) -> str:
    def replacement(match: re.Match[str]) -> str:
        kind, key = match.groups()
        values = FIGURE_NUMBERS if kind == "fig" else TABLE_NUMBERS
        label = "Figure" if kind == "fig" else "Table"
        return f"{label} {values[key]}"

    return TOKEN_PATTERN.sub(replacement, text)


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_word_field(paragraph, instruction: str, cached: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run._r.append(instruction_text)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    cached_run = paragraph.add_run(cached)
    _set_run_font(cached_run)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _bookmark_name(kind: str, key: str) -> str:
    prefix = "Fig" if kind == "fig" else "Tbl"
    return f"_Ref{prefix}{key.title().replace('_', '')}"


def _paragraph_shading(paragraph, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    p_pr.append(borders)


def _table_geometry(table, widths_inches: list[float]) -> None:
    """Set fixed Word DXA geometry: tblW, tblInd, tblGrid, and every tcW agree."""

    widths = [round(value * 1440) for value in widths_inches]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 70), ("bottom", 70), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _make_numbering(document: Document, *, decimal: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal" if decimal else "bullet")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if decimal else "\u2022")
    level.append(level_text)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "270")
    p_pr.append(indentation)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


class DocxBuilder:
    def __init__(self) -> None:
        self.document = Document()
        self.bookmark_id = 20
        self._configure_document()
        self.decimal_num_id = _make_numbering(self.document, decimal=True)

    def _configure_document(self) -> None:
        section = self.document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(17)
        section.right_margin = Mm(17)
        section.top_margin = Mm(16)
        section.bottom_margin = Mm(16)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(9)
        section.start_type = WD_SECTION.NEW_PAGE

        styles = self.document.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(BLACK)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.16
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for name, size, color, before, after in (
            ("Title", 27, NAVY, 0, 8),
            ("Heading 1", 15.5, NAVY, 8, 5),
            ("Heading 2", 12.5, TEAL, 7, 4),
            ("Heading 3", 10.5, NAVY, 5, 3),
        ):
            style = styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = name != "Title"
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        caption = styles["Caption"]
        caption.font.name = "Calibri"
        caption.font.size = Pt(7.6)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string(SLATE)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(5)
        caption.paragraph_format.keep_with_next = True
        if "Equation" not in styles:
            styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
        equation = styles["Equation"]
        equation.font.name = "Cambria Math"
        equation.font.size = Pt(10)
        equation.font.color.rgb = RGBColor.from_string(NAVY)
        equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        equation.paragraph_format.space_before = Pt(4)
        equation.paragraph_format.space_after = Pt(6)

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run("MARKETREADY FUNDS   |   FINS5545 PROJECT B")
        _set_run_font(run, size=7.5, color=SLATE, bold=True)
        footer = section.footer.paragraphs[0]
        footer.paragraph_format.space_before = Pt(0)
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("z5603162   |   ")
        _set_run_font(run, size=7.5, color=SLATE)
        _add_word_field(footer, "PAGE", "1")
        run = footer.add_run("   |   EDUCATIONAL PROTOTYPE")
        _set_run_font(run, size=7.5, color=SLATE)
        settings = self.document.settings._element
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)

        properties = self.document.core_properties
        properties.title = "MarketReady Funds - FINS5545 Project B"
        properties.subject = "Systematic multi-asset funds with news-sentiment analytics"
        properties.author = "z5603162"
        properties.keywords = "FINS5545, portfolio, sentiment, Streamlit"

    def _add_text_with_references(self, paragraph, text: str) -> None:
        position = 0
        for match in TOKEN_PATTERN.finditer(text):
            if match.start() > position:
                paragraph.add_run(text[position : match.start()])
            kind, key = match.groups()
            label = "Figure" if kind == "fig" else "Table"
            paragraph.add_run(f"{label} ")
            values = FIGURE_NUMBERS if kind == "fig" else TABLE_NUMBERS
            _add_word_field(paragraph, f"REF {_bookmark_name(kind, key)} \\h", str(values[key]))
            position = match.end()
        if position < len(text):
            paragraph.add_run(text[position:])
        for run in paragraph.runs:
            _set_run_font(run)

    def kicker(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(text.upper())
        _set_run_font(run, size=9, color=GOLD, bold=True)

    def title(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _set_run_font(run, size=27, color=NAVY, bold=True)

    def subtitle(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(text)
        _set_run_font(run, size=13, color=TEAL, bold=False)

    def meta(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(text)
        _set_run_font(run, size=8.5, color=SLATE, italic=True)

    def h1(self, text: str) -> None:
        self.document.add_heading(text, level=1)

    def h2(self, text: str) -> None:
        self.document.add_heading(text, level=2)

    def h3(self, text: str) -> None:
        self.document.add_heading(text, level=3)

    def p(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Normal")
        self._add_text_with_references(paragraph, text)

    def callout(self, label: str, text: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(4)
        paragraph.paragraph_format.right_indent = Mm(2)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(7)
        _paragraph_shading(paragraph, LIGHT, TEAL)
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, size=9.3, color=NAVY, bold=True)
        self._add_text_with_references(paragraph, text)

    def equation(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Equation")
        run = paragraph.add_run(text)
        _set_run_font(run, name="Cambria Math", size=10, color=NAVY)

    def numbered(self, items: list[tuple[str, str]]) -> None:
        for label, text in items:
            paragraph = self.document.add_paragraph(style="Normal")
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            level = OxmlElement("w:ilvl")
            level.set(qn("w:val"), "0")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), str(self.decimal_num_id))
            num_pr.append(level)
            num_pr.append(num_id)
            p_pr.append(num_pr)
            label_run = paragraph.add_run(f"{label}. ")
            _set_run_font(label_run, bold=True, color=NAVY)
            self._add_text_with_references(paragraph, text)

    def _caption(self, kind: str, key: str, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Caption")
        label = "Figure" if kind == "fig" else "Table"
        values = FIGURE_NUMBERS if kind == "fig" else TABLE_NUMBERS
        paragraph.add_run(f"{label} ")
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(self.bookmark_id))
        bookmark_start.set(qn("w:name"), _bookmark_name(kind, key))
        paragraph._p.append(bookmark_start)
        _add_word_field(paragraph, f"SEQ {label} \\* ARABIC", str(values[key]))
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(self.bookmark_id))
        paragraph._p.append(bookmark_end)
        self.bookmark_id += 1
        paragraph.add_run(f". {text}")
        for run in paragraph.runs:
            _set_run_font(run, size=7.6, color=SLATE, italic=True)

    def table(
        self,
        key: str,
        caption: str,
        headers: list[str],
        rows: list[list[str]],
        widths: list[float],
        *,
        font_size: float = 7.5,
    ) -> None:
        self._caption("table", key, caption)
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _table_geometry(table, widths)
        header = table.rows[0]
        header_pr = header._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        header_pr.append(repeat)
        for index, value in enumerate(headers):
            cell = header.cells[index]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), NAVY)
            cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(run, size=font_size, color=WHITE, bold=True)
        for row_index, values in enumerate(rows):
            row = table.add_row()
            for index, value in enumerate(values):
                cell = row.cells[index]
                cell.text = str(value)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_index % 2:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "F7FAFC")
                    cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        _set_run_font(run, size=font_size, color=BLACK)
        _table_geometry(table, widths)
        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    def figure(
        self,
        key: str,
        path: Path,
        caption: str,
        *,
        width_inches: float,
        alt_text: str,
    ) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run()
        shape = run.add_picture(str(path), width=Inches(width_inches))
        shape._inline.docPr.set("descr", alt_text)
        self._caption("fig", key, caption)

    def new_page(self) -> None:
        self.document.add_page_break()

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(path)


class SectionRule(Flowable):
    def __init__(self, width: float, color_value: colors.Color) -> None:
        super().__init__()
        self.width = width
        self.height = 5
        self.color_value = color_value

    def draw(self) -> None:
        self.canv.setStrokeColor(self.color_value)
        self.canv.setLineWidth(1.2)
        self.canv.line(0, 2, self.width, 2)


class PdfBuilder:
    def __init__(self, path: Path) -> None:
        self.path = path
        self.story: list[Flowable] = []
        self.content_width = A4[0] - 34 * rl_mm
        self.styles = self._styles()

    def _styles(self) -> dict[str, ParagraphStyle]:
        sample = getSampleStyleSheet()
        return {
            "body": ParagraphStyle(
                "MarketReady Body",
                parent=sample["BodyText"],
                fontName="Helvetica",
                fontSize=9.2,
                leading=11.2,
                textColor=rl_hex(BLACK),
                alignment=TA_JUSTIFY,
                spaceAfter=5,
            ),
            "kicker": ParagraphStyle(
                "Kicker",
                parent=sample["BodyText"],
                fontName="Helvetica-Bold",
                fontSize=9,
                leading=11,
                textColor=rl_hex(GOLD),
                alignment=TA_CENTER,
                spaceBefore=18,
                spaceAfter=10,
            ),
            "title": ParagraphStyle(
                "Title",
                parent=sample["Title"],
                fontName="Helvetica-Bold",
                fontSize=27,
                leading=31,
                textColor=rl_hex(NAVY),
                alignment=TA_CENTER,
                spaceAfter=8,
            ),
            "subtitle": ParagraphStyle(
                "Subtitle",
                parent=sample["BodyText"],
                fontName="Helvetica",
                fontSize=13,
                leading=16,
                textColor=rl_hex(TEAL),
                alignment=TA_CENTER,
                spaceAfter=10,
            ),
            "meta": ParagraphStyle(
                "Meta",
                parent=sample["BodyText"],
                fontName="Helvetica-Oblique",
                fontSize=8.5,
                leading=11,
                textColor=rl_hex(SLATE),
                alignment=TA_CENTER,
                spaceAfter=8,
            ),
            "h1": ParagraphStyle(
                "H1",
                parent=sample["Heading1"],
                fontName="Helvetica-Bold",
                fontSize=15.5,
                leading=18,
                textColor=rl_hex(NAVY),
                spaceBefore=8,
                spaceAfter=5,
                keepWithNext=True,
            ),
            "h2": ParagraphStyle(
                "H2",
                parent=sample["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=12.5,
                leading=15,
                textColor=rl_hex(TEAL),
                spaceBefore=7,
                spaceAfter=4,
                keepWithNext=True,
            ),
            "h3": ParagraphStyle(
                "H3",
                parent=sample["Heading3"],
                fontName="Helvetica-Bold",
                fontSize=10.5,
                leading=13,
                textColor=rl_hex(NAVY),
                spaceBefore=5,
                spaceAfter=3,
                keepWithNext=True,
            ),
            "caption": ParagraphStyle(
                "Caption",
                parent=sample["BodyText"],
                fontName="Helvetica-Oblique",
                fontSize=7.4,
                leading=9.2,
                textColor=rl_hex(SLATE),
                alignment=TA_LEFT,
                spaceBefore=3,
                spaceAfter=5,
                keepWithNext=True,
            ),
            "equation": ParagraphStyle(
                "Equation",
                parent=sample["BodyText"],
                fontName="Courier",
                fontSize=9.5,
                leading=12,
                textColor=rl_hex(NAVY),
                alignment=TA_CENTER,
                spaceBefore=4,
                spaceAfter=6,
            ),
            "numbered": ParagraphStyle(
                "Numbered",
                parent=sample["BodyText"],
                fontName="Helvetica",
                fontSize=9.2,
                leading=11.2,
                leftIndent=22,
                firstLineIndent=-12,
                textColor=rl_hex(BLACK),
                spaceAfter=5,
            ),
        }

    def _paragraph(self, text: str, style: str) -> Paragraph:
        safe = html.escape(_replace_static_tokens(text))
        return Paragraph(safe, self.styles[style])

    def kicker(self, text: str) -> None:
        self.story.append(self._paragraph(text.upper(), "kicker"))

    def title(self, text: str) -> None:
        self.story.append(self._paragraph(text, "title"))

    def subtitle(self, text: str) -> None:
        self.story.append(self._paragraph(text, "subtitle"))

    def meta(self, text: str) -> None:
        self.story.append(self._paragraph(text, "meta"))

    def h1(self, text: str) -> None:
        self.story.append(self._paragraph(text, "h1"))

    def h2(self, text: str) -> None:
        self.story.append(self._paragraph(text, "h2"))

    def h3(self, text: str) -> None:
        self.story.append(self._paragraph(text, "h3"))

    def p(self, text: str) -> None:
        self.story.append(self._paragraph(text, "body"))

    def callout(self, label: str, text: str) -> None:
        safe = html.escape(_replace_static_tokens(text))
        paragraph = Paragraph(
            f"<b>{html.escape(label)}:</b> {safe}",
            ParagraphStyle(
                "Callout",
                parent=self.styles["body"],
                backColor=rl_hex(LIGHT),
                borderColor=rl_hex(TEAL),
                borderWidth=0,
                borderPadding=(7, 8, 7, 10),
                leftIndent=5,
                spaceBefore=4,
                spaceAfter=7,
            ),
        )
        self.story.append(paragraph)

    def equation(self, text: str) -> None:
        self.story.append(self._paragraph(text, "equation"))

    def numbered(self, items: list[tuple[str, str]]) -> None:
        for index, (label, text) in enumerate(items, start=1):
            safe = html.escape(_replace_static_tokens(text))
            paragraph = Paragraph(
                f"{index}. <b>{html.escape(label)}.</b> {safe}", self.styles["numbered"]
            )
            self.story.append(paragraph)

    def table(
        self,
        key: str,
        caption: str,
        headers: list[str],
        rows: list[list[str]],
        widths: list[float],
        *,
        font_size: float = 7.5,
    ) -> None:
        caption_text = f"Table {TABLE_NUMBERS[key]}. {caption}"
        self.story.append(self._paragraph(caption_text, "caption"))
        data = [
            [
                Paragraph(html.escape(value), self._table_cell_style(font_size, True))
                for value in headers
            ]
        ]
        for row in rows:
            data.append(
                [
                    Paragraph(html.escape(str(value)), self._table_cell_style(font_size, False))
                    for value in row
                ]
            )
        table = Table(
            data,
            colWidths=[value * inch for value in widths],
            repeatRows=1,
            hAlign="LEFT",
        )
        style = [
            ("BACKGROUND", (0, 0), (-1, 0), rl_hex(NAVY)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("GRID", (0, 0), (-1, -1), 0.35, rl_hex("CBD5E1")),
        ]
        for row_index in range(2, len(data), 2):
            style.append(("BACKGROUND", (0, row_index), (-1, row_index), rl_hex("F7FAFC")))
        table.setStyle(TableStyle(style))
        self.story.extend([table, Spacer(1, 4)])

    def _table_cell_style(self, font_size: float, header: bool) -> ParagraphStyle:
        return ParagraphStyle(
            f"TableCell{font_size}{header}",
            fontName="Helvetica-Bold" if header else "Helvetica",
            fontSize=font_size,
            leading=font_size + 1.5,
            textColor=colors.white if header else rl_hex(BLACK),
            alignment=TA_CENTER if header else TA_LEFT,
        )

    def figure(
        self,
        key: str,
        path: Path,
        caption: str,
        *,
        width_inches: float,
        alt_text: str,
    ) -> None:
        del alt_text
        with PillowImage.open(path) as source:
            pixel_width, pixel_height = source.size
        width = width_inches * inch
        height = width * pixel_height / pixel_width
        max_height = 6.8 * inch
        if height > max_height:
            scale = max_height / height
            width *= scale
            height *= scale
        image = Image(str(path), width=width, height=height)
        caption_flowable = self._paragraph(
            f"Figure {FIGURE_NUMBERS[key]}. {caption}", "caption"
        )
        self.story.append(KeepTogether([image, caption_flowable]))

    def new_page(self) -> None:
        self.story.append(PageBreak())

    def _page(self, canvas, document) -> None:
        canvas.saveState()
        page = canvas.getPageNumber()
        width, height = A4
        if page > 1:
            canvas.setFont("Helvetica-Bold", 7.5)
            canvas.setFillColor(rl_hex(SLATE))
            canvas.drawString(17 * rl_mm, height - 10 * rl_mm, "MARKETREADY FUNDS")
            canvas.drawRightString(
                width - 17 * rl_mm, height - 10 * rl_mm, "FINS5545 PROJECT B"
            )
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(rl_hex(SLATE))
        canvas.drawString(17 * rl_mm, 8 * rl_mm, "z5603162 | EDUCATIONAL PROTOTYPE")
        canvas.drawRightString(width - 17 * rl_mm, 8 * rl_mm, f"PAGE {page}")
        canvas.restoreState()

    def build(self) -> None:
        self.path.parent.mkdir(parents=True, exist_ok=True)
        frame = Frame(
            17 * rl_mm,
            16 * rl_mm,
            A4[0] - 34 * rl_mm,
            A4[1] - 32 * rl_mm,
            leftPadding=0,
            rightPadding=0,
            topPadding=0,
            bottomPadding=0,
        )
        template = PageTemplate(id="A4 report", frames=[frame], onPage=self._page)
        document = BaseDocTemplate(
            str(self.path),
            pagesize=A4,
            leftMargin=17 * rl_mm,
            rightMargin=17 * rl_mm,
            topMargin=16 * rl_mm,
            bottomMargin=16 * rl_mm,
            title="MarketReady Funds - FINS5545 Project B",
            author="z5603162",
            subject="Systematic multi-asset funds with news-sentiment analytics",
        )
        document.addPageTemplates([template])
        document.build(self.story)


def compose_report(builder: Builder, data: ReportData) -> None:
    metrics = data.metrics.set_index("fund_id")
    equity_equal = metrics.loc["equity_equal_weight"]
    combined_rp = metrics.loc["combined_risk_parity"]
    crypto_mv = metrics.loc["crypto_min_variance"]
    base = data.fusion.set_index("strategy").loc["Base equity risk parity"]
    tilt = data.fusion.set_index("strategy").loc["Coverage-aware sentiment tilt"]
    vader = data.vader.set_index("model")
    plain = vader.loc["Plain VADER 3.3.2"]
    extended = vader.loc["Finance-extended VADER 3.3.2"]
    combined_risk = data.combined_risk.set_index("fund_id")
    combined_equal_risk = combined_risk.loc["combined_equal_weight"]
    combined_max_risk = combined_risk.loc["combined_max_sharpe"]
    combined_rp_risk = combined_risk.loc["combined_risk_parity"]
    predictive = data.sentiment_predictive.set_index("horizon_days")
    one_day = predictive.loc[1]
    five_day = predictive.loc[5]
    combined_rp_fee = data.fee_sensitivity.loc[
        data.fee_sensitivity["fund_id"].eq("combined_risk_parity")
    ].set_index("annual_management_fee")
    solver_events = data.weights[
        ["fund_id", "rebalance_date", "solver_success", "solver_message"]
    ].drop_duplicates(["fund_id", "rebalance_date"])
    fallbacks = solver_events.loc[~solver_events["solver_success"]]
    fallback_date = pd.to_datetime(fallbacks["rebalance_date"].iloc[0])
    total_optimizer_cells = int(data.imputation["optimizer_cells"].sum())
    total_mean_fills = int(data.imputation["mean_filled_cells"].sum())

    builder.kicker("FINS5545 Financial Market Data Literacy | Project B")
    builder.title("MarketReady Funds")
    builder.subtitle("Systematic multi-asset funds with coverage-aware news sentiment")
    builder.meta("z5603162 | Source sample: 1 January 2020 to 31 December 2023")
    builder.meta("Out-of-sample fund evidence: 1 January 2021 to 31 December 2023")
    builder.h1("Executive summary")
    builder.p(
        "MarketReady Funds turns the course equity, crypto, and headline datasets into "
        "12 separately investable systematic products and a decision-focused Streamlit "
        "interface. The product covers three asset families and four portfolio rules, "
        "so investors can separate a choice about market exposure from a choice about "
        "portfolio construction. All reported fund returns are walk-forward, allow "
        "holdings to drift between monthly rebalances, and deduct 5 basis points per "
        "unit of realised turnover. They are net of that estimated trading cost but "
        "before management fees, taxes, spreads, and market impact."
    )
    builder.callout(
        "Audited headline results",
        f"Crypto Minimum Variance records the highest Sharpe ratio ({crypto_mv.sharpe_ratio:.2f}) "
        f"but also a {abs(crypto_mv.max_drawdown):.1%} maximum drawdown. Combined Risk Parity "
        f"delivers a {combined_rp.sharpe_ratio:.2f} Sharpe ratio with a "
        f"{abs(combined_rp.max_drawdown):.1%} drawdown. The finance-extended VADER model changes "
        f"{extended.changed_from_plain_share:.1%} of headline scores, but the lagged signal's "
        f"one- and five-day mean rank ICs are only {one_day.mean_spearman_ic:.3f} and "
        f"{five_day.mean_spearman_ic:.3f}. Its equity tilt reduces Sharpe from "
        f"{base.sharpe_ratio:.3f} to {tilt.sharpe_ratio:.3f}, while improving maximum "
        f"drawdown by {tilt.max_drawdown - base.max_drawdown:.2%}."
    )
    builder.p(
        "The main economic result is not that optimisation always wins. Equity Equal Weight "
        f"produces a {equity_equal.sharpe_ratio:.2f} Sharpe ratio, ahead of the three optimised "
        "equity funds. Mean-sensitive Maximum Sharpe is the weakest equity rule. By contrast, "
        "the Combined Risk Parity fund balances the high upside of crypto against its large "
        "losses more effectively than the other combined products. These comparisons support "
        "a product range rather than a single claimed winner."
    )
    builder.p(
        "The sentiment extension is deliberately reported as a negative result. Its near-zero "
        "rank correlations and non-positive high-minus-low spreads do not support a profitable "
        "cross-sectional ranking in this sample; higher turnover then adds cost. The app "
        "therefore presents sentiment as an analytic with explicit coverage, not as a promise "
        "of alpha or a forecast."
    )

    builder.new_page()
    builder.h1("1. Product, investor, and data scope")
    builder.h2("The product")
    builder.p(
        "MarketReady Funds is aimed at a financially literate self-directed investor who wants "
        "rules-based diversification but needs to understand the evidence before allocating. "
        "The investable object is each fund, defined by one asset family and one portfolio "
        "method. Management could charge a fee on assets in those funds; the prototype does "
        "not process orders or money. Its value is transparent comparison: the same metrics, "
        "information clock, transaction-cost rule, and visual language apply to every product."
    )
    builder.h2("Investor journey")
    builder.numbered(
        [
            ("Compare", "Filter the 12-fund marketplace and compare risk-adjusted growth."),
            ("Inspect", "Open a fact sheet with return, risk, drawdown, and current targets."),
            ("Allocate", "Combine up to five funds and simulate a normalized historical mix."),
            ("Contextualise", "Read sector news tone beside coverage and the fusion diagnostic."),
        ]
    )
    builder.h2("Data foundation and provenance")
    builder.table(
        "inputs",
        "Inputs retained for Station 3. Dates are capped at 31 December 2023. Source: "
        "FINS5545 provided bundle and author calculations.",
        ["Input", "Retained rows", "Frequency", "Use in Part B"],
        [
            ["Equity prices", "50,300", "Equity trading days", "50 adjusted-close return series"],
            ["Crypto prices", "14,610", "Seven days per week", "10 native-calendar return series"],
            ["Unique headlines", "146,836", "Event dates", "146,830 alignable headlines scored"],
        ],
        [1.35, 1.05, 1.6, 2.75],
        font_size=7.6,
    )
    builder.p(
        "The ETL preserves genuine extreme returns and removes only key duplicates and rows "
        "outside the stated sample. Equity and crypto returns are computed inside ticker on "
        "their native calendars. Crypto returns are then left-joined to equity dates for the "
        "combined funds, intentionally excluding weekend-only moves rather than differencing "
        "price levels after a calendar merge. Six year-end headlines cannot map to a later "
        "2023 trading day, leaving 146,830 scored observations."
    )

    builder.new_page()
    builder.h1("2. Fund construction and the information clock")
    builder.p(
        "The first live weight is formed only after a full trailing estimation window. Equity "
        "and combined funds start on 4 January 2021 after 252 equity observations; crypto funds "
        "start on 1 January 2021 after 365 native-calendar observations. The first observed "
        "date of each month is a rebalance. If t is a rebalance date, every return used to "
        "estimate its target ends strictly before t."
    )
    builder.equation("target weights at t = model(returns from t - window through t - 1)")
    builder.table(
        "specifications",
        "Walk-forward fund specifications. Annualisation matches the asset calendar; caps are "
        "per asset. Source: author design.",
        ["Family", "Assets", "Window", "Annualisation", "Max weight", "Live start"],
        [
            ["Equity", "50", "252", "252", "15%", "2021-01-04"],
            ["Crypto", "10", "365", "365", "35%", "2021-01-01"],
            ["Combined", "60", "252", "252", "12%", "2021-01-04"],
        ],
        [1.15, 0.75, 0.85, 1.2, 1.05, 1.45],
        font_size=7.7,
    )
    builder.h2("Four portfolio rules")
    builder.p(
        "Equal Weight is a transparent baseline. Minimum Variance minimises forecast portfolio "
        "variance. Maximum Sharpe maximises expected excess return per unit of volatility with "
        "a zero risk-free rate. Risk Parity minimises differences between assets' variance "
        "contributions. All targets are long-only and fully invested. To reduce estimation "
        "noise, annualised means are shrunk 50% toward the cross-sectional mean and the sample "
        "covariance is shrunk 10% toward its diagonal. Assets need at least 80% of the trailing "
        "window; any remaining gaps would be filled by that asset's window mean. Combined "
        "moments use already-computed crypto daily returns observed on equity dates and are "
        "annualised at 252 periods; crypto-only moments use the full seven-day series and 365."
    )
    builder.h2("Execution and validation")
    builder.p(
        "After each monthly trade, asset weights drift with relative returns until the next "
        "rebalance. Turnover is one half of the absolute difference between the new target and "
        "the actual pre-trade weights. A cost of 5 basis points per unit of that turnover is "
        "deducted on the rebalance date. This model is more realistic than holding target "
        "weights constant every day, but it still omits spreads, market impact, taxes, fees, "
        "and capacity limits. Reported 'net' results refer only to this estimated trading-cost "
        "deduction; no management fee is included."
    )
    builder.callout(
        "Reproducibility gate",
        f"The final pipeline passes {int(data.validation['passed'].sum())} of "
        f"{len(data.validation)} checks covering unique keys, first-return boundaries, "
        "forward-only headline alignment, 12 fund identities, long-only weights summing to one, "
        "strict estimation timing, method differentiation, solver reliability, sentiment bounds, "
        "signal-calendar coverage, risk-contribution reconciliation, annual outputs, and fusion "
        f"weights. Across {total_optimizer_cells:,} repeated estimation-window cells, "
        f"{total_mean_fills:,} are mean-filled and no asset-window fails the 80% coverage rule. "
        f"There are {len(fallbacks)} solver fallback event(s). On failure or a non-finite SLSQP "
        "solution, the same cap is applied to inverse variance for Minimum Variance, positive "
        "mean divided by variance for Maximum Sharpe, or inverse volatility for Risk Parity."
    )

    builder.new_page()
    builder.h1("3. Out-of-sample fund comparison")
    builder.table(
        "performance",
        "Out-of-sample fund performance after 5 basis points per unit of realised turnover but "
        "before management fees. Return is the annualised arithmetic mean; volatility "
        "uses the matching 252- or 365-period calendar; Sharpe assumes a 0% risk-free rate; "
        "turnover is average realised monthly turnover. Source sample: 2020-2023; live period: "
        "2021-2023. Source: author calculations.",
        ["Fund", "Return", "Vol.", "Sharpe", "Max DD", "Turnover"],
        performance_rows(data.metrics),
        [2.45, 0.82, 0.78, 0.72, 0.82, 0.9],
        font_size=7.0,
    )
    builder.p(
        "{table:performance} separates high return from "
        "usable risk. Crypto Minimum Variance leads the Sharpe ranking at "
        f"{crypto_mv.sharpe_ratio:.2f} and compounds $1 to ${crypto_mv.terminal_wealth:.2f}, but "
        f"its {abs(crypto_mv.max_drawdown):.1%} maximum drawdown is outside a conventional "
        "balanced-investor tolerance. Crypto therefore supplies the largest upside and the "
        "largest path risk. Annualised arithmetic returns should not be mistaken for CAGRs; "
        "volatility drag makes that distinction material for crypto."
    )
    builder.p(
        f"Combined Risk Parity is the strongest combined product, with a "
        f"{combined_rp.sharpe_ratio:.2f} Sharpe ratio, {pct(float(combined_rp.annualized_return))} "
        f"annualised return, and {pct(float(combined_rp.max_drawdown))} maximum drawdown. It "
        "outperforms Combined Equal Weight on Sharpe because its volatility-based targets keep "
        "crypto near a single-digit to low-teens family weight. Combined Minimum Variance "
        "suppresses volatility further but gives up too much return over this live sample. A "
        f"1.0% annual management fee would reduce Combined Risk Parity's reported return to "
        f"{combined_rp_fee.loc[0.01, 'annualized_return']:.1%} and Sharpe to "
        f"{combined_rp_fee.loc[0.01, 'sharpe_ratio']:.2f}; it remains the highest-Sharpe "
        "combined fund under that simple fee sensitivity."
    )
    builder.p(
        f"Within equities, the simple Equal Weight fund has the highest Sharpe ratio "
        f"({equity_equal.sharpe_ratio:.2f}). Maximum Sharpe ranks last at "
        f"{metrics.loc['equity_max_sharpe', 'sharpe_ratio']:.2f} and carries the highest average "
        "equity turnover. The result is consistent with noisy mean estimates: optimising a "
        "ratio built on expected returns can amplify sampling error even after shrinkage. The "
        "comparison argues for showing several rules and their costs rather than presenting "
        "the in-sample objective as an investor guarantee."
    )

    builder.new_page()
    builder.h1("4. Growth paths: asset family dominates method")
    builder.figure(
        "growth",
        RESULTS_DIR / "figures" / "growth_of_dollar.png",
        "Growth of $1 for four monthly walk-forward methods in equity, crypto, and combined "
        "families, after 5 basis points per unit of realised turnover and before management "
        "fees. Log scales allow "
        "proportional comparison. Live period: 2021-2023. Source: FINS5545 data and author "
        "calculations.",
        width_inches=6.65,
        alt_text="Three panels comparing growth of one dollar for 12 systematic funds.",
    )
    builder.p(
        "{fig:growth} shows that the choice of asset family creates a wider range of outcomes "
        "than the choice of optimiser. Equity and combined products finish in a relatively "
        "narrow band, while crypto paths rise rapidly in 2021, lose most of those gains through "
        "2022, and only partly recover in 2023. A terminal value alone would hide that sequence."
    )
    builder.p(
        "Minimum Variance is the clear crypto leader at the sample end, but it still spends a "
        "long period more than 60% below its prior peak. Within equities, Equal Weight and Risk "
        "Parity recover more strongly than the two moment-optimised alternatives. Within the "
        "combined family, Risk Parity offers the best return-risk compromise, while Equal Weight "
        "accepts the larger crypto allocation mechanically created by one weight per asset."
    )
    builder.p(
        "Annual results in {table:annual} make the regime dependence explicit. Combined Risk "
        "Parity returns 37.3% in 2021, -6.8% in 2022, and 15.5% in 2023; Crypto Minimum Variance "
        "moves from 328.1% to -59.3% and then 147.2%. Method rankings vary by year, but the core "
        "conclusions remain: crypto drives the widest outcomes, Risk Parity limits its combined-"
        "fund influence, and no three-year Sharpe ranking is a forecast of the next regime."
    )

    builder.new_page()
    builder.h1("5. Downside experience is part of the product")
    builder.figure(
        "drawdown",
        RESULTS_DIR / "figures" / "drawdown.png",
        "Drawdown from the prior wealth peak for the highest-Sharpe fund in each family. "
        "Live period: 2021-2023. Source: FINS5545 data and author calculations.",
        width_inches=5.15,
        alt_text=(
            "Drawdown paths for Combined Risk Parity, Crypto Minimum Variance, and "
            "Equity Equal Weight."
        ),
    )
    builder.p(
        "{fig:drawdown} changes the product conclusion. Combined Risk Parity and Equity Equal "
        "Weight both lose roughly one fifth from peak to trough, while Crypto Minimum Variance "
        "falls more than 70%. Crypto's Sharpe ratio above one does not make its loss experience "
        "moderate. The high annualised mean and extreme drawdown coexist because the series is "
        "highly volatile and path dependent."
    )
    builder.p(
        "For the target user, maximum drawdown is easier to translate into a funding decision "
        "than volatility alone. A 72% loss requires a subsequent gain of about 257% to recover. "
        "That arithmetic is why MarketReady Funds exposes drawdown in every fact sheet and why "
        "the allocation lab does not default to the single highest-Sharpe fund. A fund range "
        "should let investors express different loss tolerances rather than collapse the choice "
        "to a leaderboard."
    )

    builder.new_page()
    builder.h1("6. Targets, drift, and fact sheets")
    builder.figure(
        "weights",
        RESULTS_DIR / "figures" / "weights_over_time.png",
        "Target equity and crypto family weights for the four combined funds at monthly "
        "rebalances. Values between rebalances drift with returns and are not shown. Live "
        "period: 2021-2023. Source: author calculations.",
        width_inches=5.7,
        alt_text=(
            "Four panels showing equity and crypto target weights through time for "
            "combined funds."
        ),
    )
    builder.p(
        "{fig:weights} shows capital allocation, but capital weight is not risk weight. The "
        "diagnostic below uses each rebalance's same 252-period, diagonally shrunk combined "
        "covariance matrix to decompose target variance into equity and crypto components."
    )
    builder.table(
        "combined_risk",
        "Combined-fund target diagnostics across 36 monthly rebalances. Capital and covariance-"
        "risk figures are medians; the capital range is in parentheses. Effective holdings are "
        "1/HHI. Source: author calculations.",
        ["Method", "Crypto capital", "Crypto risk", "Effective N", "Top 5", "Turnover"],
        combined_risk_rows(data),
        [1.25, 1.55, 1.0, 0.9, 0.8, 0.9],
        font_size=6.9,
    )
    builder.p(
        "{table:combined_risk} explains why the methods behave differently. Equal Weight puts "
        f"{combined_equal_risk.median_crypto_capital_weight:.1%} of capital in crypto but a "
        f"median {combined_equal_risk.median_crypto_risk_contribution:.1%} of estimated variance "
        "comes from crypto. Risk Parity uses median crypto capital of only "
        f"{combined_rp_risk.median_crypto_capital_weight:.1%} and holds its family risk share at "
        f"{combined_rp_risk.median_crypto_risk_contribution:.1%}: 10 of 60 equal asset-level risk "
        "budgets. Maximum Sharpe is visibly less stable: crypto capital ranges from "
        f"{combined_max_risk.minimum_crypto_capital_weight:.1%} to "
        f"{combined_max_risk.maximum_crypto_capital_weight:.1%}, its median effective holdings "
        f"fall to {combined_max_risk.median_effective_holdings:.1f}, its top five positions hold "
        f"{combined_max_risk.median_top_5_weight:.1%}, and average turnover is "
        f"{metrics.loc['combined_max_sharpe', 'average_turnover']:.1%}. These are sample "
        "diagnostics of mean-sensitive concentration, not forecasts of future instability."
    )
    builder.h2("Twelve consistent fact sheets")
    builder.p(
        "Each app fact sheet joins the fund's live growth and drawdown path to annualised return, "
        "annualised volatility, Sharpe ratio, maximum drawdown, ending value, and the complete "
        "latest target portfolio. Combined fact sheets also show capital versus covariance-risk "
        "share and concentration. The latest three largest holdings for every fund are audited "
        "in {table:holdings}. Full targets remain downloadable from the app and in "
        "results/data/fund_weights.csv. This design keeps the report comparative while letting "
        "the product serve position-level detail on demand."
    )

    builder.new_page()
    builder.h1("7. Sector sentiment: tone must be read beside coverage")
    builder.p(
        "VADER 3.3.2 scores each original headline without lowercasing or removing punctuation, "
        "negation, boosters, or capitalisation. Those features are inputs to the model, not "
        "formatting noise. Its compound score bounds the adjusted valence sum x as follows "
        "(Hutto and Gilbert, 2014):"
    )
    builder.equation("compound = x / sqrt(x^2 + 15)")
    builder.p(
        "The finance extension adds 20 reviewed words, 9 reviewed phrases, and 5 intensity "
        "modifiers from the Week 8 ten-rater exercise. The code snapshots and restores VADER's "
        "module dictionaries so repeated runs do not contaminate the baseline. Headline scores "
        "are averaged within ticker-day first; observed ticker-day means are then equally "
        "weighted within sector. No-news observations remain missing in the headline index, "
        "while a separate neutral-fill series multiplies observed tone by ticker coverage."
    )
    builder.p(
        f"The finance extension changes {extended.changed_from_plain_share:.1%} of scores and "
        f"reduces the share classified as neutral from {plain.neutral_share:.1%} to "
        f"{extended.neutral_share:.1%}. {{table:vader}} reports the full distribution. This is "
        "dictionary-coverage evidence, not an accuracy claim, because no labelled finance-"
        "headline test set is available."
    )
    builder.figure(
        "sentiment",
        RESULTS_DIR / "figures" / "sector_sentiment_index.png",
        "Twenty-one-trading-day mean finance-extended VADER compound score by equity sector, "
        "split by average ticker-news coverage. Index dates: 2020-2023. Source: FINS5545 "
        "headlines and author calculations.",
        width_inches=5.25,
        alt_text=(
            "Two panels of sector sentiment paths split into higher- and "
            "lower-coverage sectors."
        ),
    )
    builder.table(
        "sentiment_predictive",
        "Cross-sectional validation of the effective lagged ticker signal against future equity "
        "returns. IC is the daily Spearman rank correlation; H-L is the equal-weight top-minus-"
        "bottom sentiment quintile return in basis points. HAC t-statistics use horizon minus "
        "one lag. Signal dates match the live equity fund period, 2021-2023. Source: author "
        "calculations.",
        ["Horizon", "Mean IC", "IC t", "H-L (bp)", "H-L t", "IC > 0"],
        sentiment_predictive_rows(data),
        [1.0, 0.9, 0.8, 1.0, 0.8, 0.9],
        font_size=7.0,
    )
    builder.p(
        "{fig:sentiment} is therefore a descriptive tone index, not evidence of return "
        "predictability. {table:sentiment_predictive} finds a one-day mean IC of "
        f"{one_day.mean_spearman_ic:.3f} (HAC t={one_day.ic_hac_t_stat:.2f}) and a five-day IC "
        f"of {five_day.mean_spearman_ic:.3f} (t={five_day.ic_hac_t_stat:.2f}). The high-minus-low "
        f"spread is {one_day.mean_high_minus_low_return * 10_000:+.2f} bp at one day and "
        f"{five_day.mean_high_minus_low_return * 10_000:+.2f} bp at five days. These tests do not "
        "show an economically or statistically persuasive positive ranking relation in this "
        "sample; the five-day observations overlap, so its HAC result remains descriptive."
    )

    builder.new_page()
    builder.h1("8. Fusion extension: an informative negative result")
    builder.p(
        "The fusion starts from Equity Risk Parity. For each ticker, the model forms a "
        "21-trading-day mean headline score, lags both tone and coverage by one trading day, "
        "standardises available tone across tickers, and scales the z-score by the square root "
        "of lagged coverage. At each monthly rebalance, the base target is multiplied by "
        "exp(0.30 times the effective signal) and projected back to the long-only 15% cap. "
        "Saturday and Monday headlines aligned to Monday can therefore first influence Tuesday; "
        "a Monday rebalance never uses Monday coverage."
    )
    builder.table(
        "fusion",
        "Before-and-after equity sentiment fusion after 5 basis points per unit of realised "
        "turnover and before management fees. Live period: 2021-2023. Source: author calculations.",
        ["Strategy", "Return", "Vol.", "Sharpe", "Max DD", "Turnover"],
        [
            [
                "Base risk parity",
                pct(float(base.annualized_return)),
                pct(float(base.annualized_volatility)),
                number(float(base.sharpe_ratio), 3),
                pct(float(base.max_drawdown)),
                pct(float(base.average_turnover)),
            ],
            [
                "Coverage-aware tilt",
                pct(float(tilt.annualized_return)),
                pct(float(tilt.annualized_volatility)),
                number(float(tilt.sharpe_ratio), 3),
                pct(float(tilt.max_drawdown)),
                pct(float(tilt.average_turnover)),
            ],
        ],
        [1.65, 0.85, 0.85, 0.8, 0.85, 1.0],
        font_size=7.5,
    )
    builder.figure(
        "fusion",
        RESULTS_DIR / "figures" / "fusion_comparison.png",
        "Growth and drawdown for Equity Risk Parity before and after the one-day-lagged, "
        "coverage-aware sentiment tilt. Both strategies include turnover costs. Live period: "
        "2021-2023. Source: author calculations.",
        width_inches=6.2,
        alt_text="Two panels comparing base equity risk parity with its sentiment-tilted version.",
    )
    builder.p(
        "{table:fusion} and {fig:fusion} reject a simple alpha story in this sample. The tilt "
        "lowers annualised "
        f"return by {base.annualized_return - tilt.annualized_return:.2%}, lowers Sharpe by "
        f"{base.sharpe_ratio - tilt.sharpe_ratio:.3f}, and raises average monthly turnover from "
        f"{base.average_turnover:.1%} to {tilt.average_turnover:.1%}. Maximum drawdown improves by "
        f"{tilt.max_drawdown - base.max_drawdown:.2%}. Before trading costs, the return gap is "
        f"{base.gross_annualized_return - tilt.gross_annualized_return:.2%}; the tilt's "
        "incremental "
        f"annual trading-cost drag is only "
        f"{tilt.annualized_trading_cost_drag - base.annualized_trading_cost_drag:.2%}. The near-"
        "zero ICs and non-positive spreads therefore explain more of the underperformance than "
        "cost alone; turnover modestly worsens it."
    )
    builder.p(
        "The extension still adds product value: it demonstrates a complete, auditable use of "
        "unstructured data with an explicit information clock, a reviewed domain lexicon, and a "
        "coverage control. {table:fusion_sensitivity} shows the conclusion is qualitatively "
        "stable: a half-strength 21-day tilt and a slower 42-day tilt also remain below the "
        f"base Sharpe of {base.sharpe_ratio:.3f}. This is not a parameter search for a winner; "
        "the two alternatives are prespecified implementation checks. The app keeps sentiment "
        "as an analytic until labelled text validation and longer return evidence support "
        "promotion."
    )

    builder.new_page()
    builder.h1("9. App implementation, reflection, and recommendations")
    builder.h2("A lightweight deployment architecture")
    builder.p(
        "The root Streamlit entrypoint reads committed CSV and JSON artifacts only. It does not "
        "download raw data, run VADER, solve an optimisation, or backtest on a viewer's machine. "
        "A cached loader validates required columns and dates before any page renders. Pure data "
        "helpers calculate current holdings and custom fund-level allocations; chart helpers "
        "apply one Plotly design system. URL query parameters preserve the active view, and every "
        "investor-facing table or simulated path can be downloaded. Combined fact sheets expose "
        "capital versus covariance-risk shares; the allocation lab aggregates latest targets to "
        "show equity-versus-crypto capital exposure, major equity sectors, top underlying "
        "holdings, effective holdings, and top-five concentration. All five views pass headless "
        "AppTest. The app is live on Streamlit Community Cloud and deployed from the main branch "
        "of the public GitHub repository."
    )
    builder.h2("What worked and what did not")
    builder.p(
        "The strongest fund result is diversification by risk rather than estimation of means. "
        "Combined Risk Parity converts a median 7.6% crypto capital allocation into a 16.7% "
        "crypto covariance-risk budget and a higher Sharpe ratio than the equity risk-parity "
        "base without inheriting crypto's full drawdown. Equal Weight is "
        "also difficult to beat within equities. These outcomes favour transparent rules and "
        "show why optimised objectives must be judged out of sample."
    )
    builder.p(
        "The finance lexicon reduces the share classified as neutral, but the project cannot "
        "claim better classification without labelled finance headlines. Coverage remains "
        "unequal: average ticker coverage is about 94% in Consumer and 53% in Materials. The "
        "return tests then find near-zero rank ICs and no positive high-minus-low spread. Three "
        "live fund years are also too short to distinguish a persistent relationship from one "
        "market regime."
    )
    builder.h2("Three real-world recommendations")
    builder.numbered(
        [
            (
                "Keep Combined Risk Parity as the balanced flagship",
                "Position it for diversified-growth investors able to tolerate a roughly 20% "
                "historical drawdown. Explain that 7.6% median crypto capital supplies 16.7% of "
                "estimated risk; do not default to Crypto Minimum Variance despite its Sharpe.",
            ),
            (
                "Hold the sentiment tilt in research",
                "Require a labelled headline validation sample, publisher and duplicate-event "
                "controls, and rolling parameter tests before any client capital follows it.",
            ),
            (
                "Add implementation realism before accepting client capital",
                "Quote trading-cost-adjusted and management-fee-adjusted returns separately, "
                "publish look-through overlap and concentration, and stress spreads, slippage, "
                "capacity, and rebalance timing before accepting client capital.",
            ),
        ]
    )
    builder.h2("Conclusion")
    builder.p(
        "MarketReady Funds completes all four Data Factory Floor stages: a verified foundation, "
        "return and text features, walk-forward funds plus sentiment analytics, and a tested "
        "investor interface. The project produces a credible product range because it makes its "
        "negative results as visible as its strong ones. Twelve fact sheets, explicit timing, "
        "audited coverage, reproducible artifacts, and a lightweight deployed app are available "
        "through the live service and its public source repository."
    )

    builder.new_page()
    builder.h1("References")
    builder.p(
        "FINS5545 Course Team. (2026). FinTech Project 2026 - FINS5545: Financial Market Data "
        "Literacy [Project brief and provided market/news data bundle]. UNSW Business School."
    )
    builder.p(
        "Hutto, C. J., and Gilbert, E. (2014). VADER: A parsimonious rule-based model for "
        "sentiment analysis of social media text. Proceedings of the International AAAI "
        "Conference on Web and Social Media, 8(1), 216-225. "
        "https://doi.org/10.1609/icwsm.v8i1.14550"
    )
    builder.h2("Report basis")
    builder.p(
        "All numerical statements and exhibits in this report are generated from the committed "
        "results/data and results/tables artifacts. The reproducible build command is "
        "python scripts/run_part_b.py, followed by python scripts/build_report.py. The Word file "
        "is the editable source; the PDF is generated from the same content model."
    )
    builder.h2("Project links")
    builder.p(f"Live app: {LIVE_APP_URL}")
    builder.p(f"Public GitHub repository: {PUBLIC_GITHUB_URL}")
    builder.p(
        "The live app is deployed from the repository's main branch with streamlit_app.py at "
        "the repository root."
    )

    builder.new_page()
    builder.h1("Appendix A. Risk-adjusted ranking")
    builder.figure(
        "sharpe",
        RESULTS_DIR / "figures" / "sharpe_barplot.png",
        "Annualised Sharpe ratio for all 12 funds after 5 basis points per unit of realised "
        "turnover and before management fees, using a 0% risk-free rate. Live period: "
        "2021-2023. Source: FINS5545 data and author calculations.",
        width_inches=5.7,
        alt_text="Horizontal bars ranking the Sharpe ratio of all 12 funds.",
    )
    builder.p(
        "{fig:sharpe} confirms that no single optimisation method dominates across families. "
        "Minimum Variance leads crypto, Risk Parity leads combined funds, and Equal Weight leads "
        "equities. The ranking is consistent with the report's recommendation to sell a range "
        "of clearly differentiated products rather than one universal rule."
    )

    builder.new_page()
    builder.h1("Appendix B. Latest holdings summary")
    builder.table(
        "holdings",
        "Three largest target positions for every fund at its latest 2023 rebalance. Full "
        "holdings and exact weights are in results/data/fund_weights.csv and the app fact sheets. "
        "Source: author calculations.",
        ["Fund", "Target date", "Three largest targets"],
        top_holdings_rows(data),
        [2.35, 1.05, 3.35],
        font_size=7.1,
    )

    builder.new_page()
    builder.h1("Appendix C. Validation and source integrity")
    builder.table(
        "validation",
        "Critical pipeline validation results. Every check must pass before artifacts are "
        "written. Source: automated author tests.",
        ["Validation check", "Passed"],
        [
            [str(row.check), "Yes" if bool(row.passed) else "No"]
            for row in data.validation.itertuples()
        ],
        [5.75, 1.0],
        font_size=7.0,
    )
    builder.table(
        "integrity",
        "Retained source-integrity checks after the mandatory sample cap and exact-key "
        "deduplication. Source: FINS5545 data and author calculations.",
        ["Check", "Count", "Treatment"],
        [
            [str(row.check), f"{int(row.count):,}", str(row.treatment)]
            for row in data.integrity.itertuples()
        ],
        [2.45, 0.85, 3.45],
        font_size=6.9,
    )

    builder.new_page()
    builder.h1("Appendix D. Annual results")
    builder.table(
        "annual",
        "Calendar-year total returns after the 5-basis-point trading-cost model and before "
        "management fees. Partial-year observations are not present; every fund has a full "
        "2021, 2022, and 2023 live year. Source: author calculations.",
        ["Fund", "2021", "2022", "2023"],
        annual_return_rows(data),
        [3.55, 1.0, 1.0, 1.0],
        font_size=7.1,
    )
    builder.p(
        "{table:annual} shows that no method wins every calendar year. The three-year product "
        "conclusions nevertheless survive the year split: crypto remains the widest-upside and "
        "deepest-downside family, while Combined Risk Parity participates in 2021 and 2023 gains "
        "with a materially smaller 2022 loss than Combined Equal Weight."
    )

    builder.new_page()
    builder.h1("Appendix E. Assumption robustness")
    builder.table(
        "cost_fee",
        "Combined Risk Parity sensitivity. Trading-cost scenarios reprice realised turnover; "
        "management fees are charged continuously to the baseline 5-basis-point path. Returns "
        "are annualised arithmetic means. Live period: 2021-2023. Source: author calculations.",
        ["Scenario", "Return", "Sharpe", "Max DD", "Ending $1"],
        cost_fee_rows(data),
        [2.65, 0.9, 0.8, 0.9, 0.9],
        font_size=7.0,
    )
    builder.table(
        "fusion_sensitivity",
        "Prespecified sentiment-fusion implementation checks, all after 5 basis points per unit "
        "of realised turnover and before management fees. These are robustness checks, not a "
        "parameter search. Live period: 2021-2023. Source: author calculations.",
        ["Scenario", "Return", "Sharpe", "Max DD", "Turnover"],
        fusion_sensitivity_rows(data),
        [2.7, 0.85, 0.8, 0.85, 0.9],
        font_size=6.9,
    )
    builder.p(
        "{table:cost_fee} shows that doubling the trading-cost rate does not alter the combined-"
        "fund recommendation, while a 1% annual management fee reduces but does not reverse the "
        "flagship's Sharpe advantage. {table:fusion_sensitivity} shows every prespecified tilt "
        "variant below the equity Risk Parity base, so the research-only sentiment conclusion "
        "is qualitatively stable."
    )

    builder.new_page()
    builder.h1("Appendix F. Text model and estimation audit")
    builder.table(
        "vader",
        "Plain and finance-extended VADER distributions across 146,830 aligned headlines. "
        "Thresholds are positive >= 0.05, negative <= -0.05, and neutral otherwise. Source: "
        "vaderSentiment 3.3.2 and author extension.",
        ["Model", "Mean", "Positive", "Neutral", "Negative", "Changed"],
        [
            [
                "Plain VADER",
                f"{plain.mean_compound:+.3f}",
                pct(float(plain.positive_share)),
                pct(float(plain.neutral_share)),
                pct(float(plain.negative_share)),
                "0.0%",
            ],
            [
                "Finance-extended",
                f"{extended.mean_compound:+.3f}",
                pct(float(extended.positive_share)),
                pct(float(extended.neutral_share)),
                pct(float(extended.negative_share)),
                pct(float(extended.changed_from_plain_share)),
            ],
        ],
        [1.55, 0.85, 1.05, 1.0, 1.0, 1.0],
        font_size=7.5,
    )
    builder.table(
        "imputation",
        "Return-data availability across repeated monthly estimation windows. Optimizer cells "
        "count overlapping window observations; exclusions count an asset-window failing the "
        "80% rule. Source: automated author audit.",
        ["Family", "Windows", "Cells", "Mean-filled", "Share", "Exclusions"],
        imputation_rows(data),
        [1.15, 0.85, 1.25, 1.15, 0.9, 1.05],
        font_size=7.1,
    )
    builder.p(
        "{table:imputation} records zero return imputations and zero asset-window exclusions "
        f"across {total_optimizer_cells:,} optimizer cells. The single fallback in "
        f"Combined Maximum Sharpe on {fallback_date.day} {fallback_date:%B %Y} is therefore "
        "a solver event, not "
        "missing-data repair. Its "
        "positive-mean/variance fallback is normalized under the same long-only 12% cap."
    )


def structural_docx_audit(path: Path) -> dict[str, int]:
    """Check the editable source without relying on a Word renderer."""

    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    combined = "\n".join(paragraphs)
    forbidden = (
        "TODO",
        "PLACEHOLDER",
        "HUMAN EDIT REQUIRED",
        ":codex-file-citation",
        "deployment have not yet",
        "deployment stage",
    )
    if any(token in combined for token in forbidden):
        raise AssertionError("report.docx contains an unresolved placeholder")
    if LIVE_APP_URL not in combined or PUBLIC_GITHUB_URL not in combined:
        raise AssertionError("report.docx is missing the final project links")
    headings = sum(paragraph.style.name.startswith("Heading") for paragraph in document.paragraphs)
    captions = sum(paragraph.style.name == "Caption" for paragraph in document.paragraphs)
    if headings < 10 or captions != len(FIGURE_NUMBERS) + len(TABLE_NUMBERS):
        raise AssertionError("report.docx structure is incomplete")
    if len(document.inline_shapes) != len(FIGURE_NUMBERS):
        raise AssertionError("report.docx figure count is incomplete")
    if len(document.tables) != len(TABLE_NUMBERS):
        raise AssertionError("report.docx table count is incomplete")
    return {
        "paragraphs": len(document.paragraphs),
        "headings": headings,
        "captions": captions,
        "tables": len(document.tables),
        "figures": len(document.inline_shapes),
    }


def pdf_audit(path: Path, render_dir: Path) -> dict[str, int]:
    """Validate and render every PDF page for visual inspection."""

    reader = PdfReader(path)
    if len(reader.pages) != 17:
        raise AssertionError(
            "report.pdf must retain 10 narrative pages plus 7 reference/appendix pages"
        )
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    forbidden = (
        "TODO",
        "PLACEHOLDER",
        "HUMAN EDIT REQUIRED",
        ":codex-file-citation",
        "deployment have not yet",
        "deployment stage",
    )
    if any(token in text for token in forbidden):
        raise AssertionError("report.pdf contains an unresolved placeholder")
    if LIVE_APP_URL not in text or PUBLIC_GITHUB_URL not in text:
        raise AssertionError("report.pdf is missing the final project links")
    with pdfplumber.open(path) as pdf:
        if len(pdf.pages) != len(reader.pages):
            raise AssertionError("PDF readers disagree on the page count")
        if not (pdf.pages[0].extract_text() or "").startswith("FINS5545"):
            raise AssertionError("PDF title page text is missing")

    if render_dir.exists():
        shutil.rmtree(render_dir)
    render_dir.mkdir(parents=True)
    document = pymupdf.open(path)
    for index, page in enumerate(document):
        pixmap = page.get_pixmap(matrix=pymupdf.Matrix(1.7, 1.7), alpha=False)
        pixmap.save(render_dir / f"page-{index + 1:02d}.png")
    return {"pages": len(reader.pages), "rendered_pages": len(list(render_dir.glob("*.png")))}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--render-dir",
        type=Path,
        default=Path("/private/tmp/z5603162_projectB_report_qa"),
    )
    args = parser.parse_args()
    data = load_report_data()
    docx_builder = DocxBuilder()
    compose_report(docx_builder, data)
    docx_builder.save(DOCX_PATH)
    pdf_builder = PdfBuilder(PDF_PATH)
    compose_report(pdf_builder, data)
    pdf_builder.build()
    docx_result = structural_docx_audit(DOCX_PATH)
    pdf_result = pdf_audit(PDF_PATH, args.render_dir)
    print(f"Created {DOCX_PATH.relative_to(PROJECT_ROOT)}: {docx_result}")
    print(f"Created {PDF_PATH.relative_to(PROJECT_ROOT)}: {pdf_result}")
    print(f"Rendered PDF pages for QA: {args.render_dir}")


if __name__ == "__main__":
    main()
